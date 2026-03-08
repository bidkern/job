from pathlib import Path
import re
import zipfile
from xml.etree import ElementTree as ET


def _extract_pdf_text_with_ocr(path: Path) -> str:
    try:
        import pypdfium2 as pdfium  # type: ignore
        from rapidocr_onnxruntime import RapidOCR  # type: ignore
    except Exception:
        return ""

    try:
        doc = pdfium.PdfDocument(str(path))
        ocr = RapidOCR()
        chunks: list[str] = []
        for idx in range(len(doc)):
            page = doc[idx]
            bitmap = page.render(scale=2.0)
            image = bitmap.to_pil()
            result, _ = ocr(image)
            if not result:
                continue
            lines = [entry[1] for entry in result if entry and len(entry) > 1 and entry[1]]
            if lines:
                chunks.append("\n".join(lines))
        return "\n".join(chunks).strip()
    except Exception:
        return ""


def _extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception:
        return _extract_pdf_text_with_ocr(path)

    try:
        reader = PdfReader(str(path))
        text = "\n".join((page.extract_text() or "") for page in reader.pages).strip()
        if text:
            return text
        return _extract_pdf_text_with_ocr(path)
    except Exception:
        return _extract_pdf_text_with_ocr(path)


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return _extract_docx_text_fallback(path)

    try:
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs).strip()
        return text or _extract_docx_text_fallback(path)
    except Exception:
        return _extract_docx_text_fallback(path)


def _extract_docx_text_fallback(path: Path) -> str:
    try:
        with zipfile.ZipFile(path, "r") as archive:
            members = [
                name
                for name in archive.namelist()
                if re.match(r"^word/(document|header\d+|footer\d+)\.xml$", name)
            ]
            xml_chunks = [archive.read(name) for name in members]
    except Exception:
        return ""

    paragraphs: list[str] = []
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for raw_xml in xml_chunks:
        try:
            root = ET.fromstring(raw_xml)
        except Exception:
            continue
        for paragraph in root.findall(".//w:p", ns):
            parts: list[str] = []
            for node in paragraph.findall(".//w:t", ns):
                if node.text:
                    parts.append(node.text)
            if parts:
                paragraphs.append("".join(parts))
    return "\n".join(paragraphs).strip()


def extract_resume_text(path: str | Path) -> str:
    target = Path(path)
    if not target.exists():
        return ""

    suffix = target.suffix.lower()
    if suffix == ".txt":
        try:
            return target.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""
    if suffix == ".pdf":
        return _extract_pdf_text(target)
    if suffix == ".docx":
        return _extract_docx_text(target)
    if suffix == ".doc":
        try:
            return target.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""
    return ""
